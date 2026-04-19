"""
CV Modernizer — FastAPI Backend
================================
Flow for /api/extract:
  1. Read file bytes in-memory (never touch disk)
  2. Extract raw text (PyMuPDF for PDF, python-docx for DOCX)
  3. Guard: reject if no email OR phone found in raw text
  4. Guard: check Supabase — reject with retry-time if same email (hashed)
             was processed within the last 3 hours (TTL on updated_at)
  5. Parse full resume with Gemini 1.5 Flash → structured JSON
  6. Store minimal record to Supabase (hashed email, hashed phone,
     experience 1-liners, education — NO bullets, NO skills, NO names)
  7. Return full resume_data JSON to frontend
"""

import os
import json
import re
import logging
import hashlib
from datetime import datetime, timedelta, timezone
from typing import List, Optional

from fastapi import FastAPI, UploadFile, File, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv

from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

load_dotenv()

# ─── Logging ────────────────────────────────────────────────────────────────
logging.basicConfig(level=logging.INFO, format="%(levelname)s  %(message)s")
log = logging.getLogger("cv_modernizer")

# ─── FastAPI App ─────────────────────────────────────────────────────────────
app = FastAPI(title="CV Modernizer API", version="1.1.0")

import os
allowed_origins_str = os.getenv("ALLOWED_ORIGINS", "http://localhost:3000,http://localhost:3001")
origins = [o.strip() for o in allowed_origins_str.split(",") if o.strip()]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# ─── Regex Patterns ──────────────────────────────────────────────────────────
# Matches standard email addresses
EMAIL_RE = re.compile(
    r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}",
    re.IGNORECASE,
)
# Matches common international phone formats (at least 7 digits)
PHONE_RE = re.compile(
    r"(?:\+?\d{1,3}[\s\-.]?)?"           # optional country code
    r"\(?\d{2,4}\)?"                      # optional area code
    r"[\s\-.]?\d{2,5}"                    # first block
    r"(?:[\s\-.]?\d{1,5}){1,3}",          # additional blocks
)

TTL_HOURS = 3

MAX_FILE_BYTES = 5 * 1024 * 1024   # 5 MB  — reject files larger than this
MAX_TEXT_CHARS = 15_000             # ~3,000 words — truncate before sending to Gemini

# Model to use — override via GEMINI_MODEL in .env
# gemini-1.5-flash-8b: lighter, higher free-tier quota, perfect for extraction
# gemini-1.5-flash-002: full 1.5 flash if you need more capability
GEMINI_MODEL = os.environ.get("GEMINI_MODEL", "gemini-1.5-flash-8b")

# ─── Supabase Client (lazy init) ─────────────────────────────────────────────
_supabase = None

def get_supabase():
    global _supabase
    if _supabase is None:
        url = os.environ.get("SUPABASE_URL", "").strip()
        key = os.environ.get("SUPABASE_KEY", "").strip()
        if not url or not key:
            raise HTTPException(
                status_code=503,
                detail=(
                    "Supabase is not configured. "
                    "Add SUPABASE_URL and SUPABASE_KEY to your .env file."
                ),
            )
        from supabase import create_client
        _supabase = create_client(url, key)
    return _supabase


# ─── Pydantic Models ─────────────────────────────────────────────────────────
class ResumeData(BaseModel):
    bio: dict
    links: List[dict]
    experience: List[dict]
    education: List[dict]
    skills: List[dict]

class GapAnalysisRequest(BaseModel):
    resume_data: ResumeData
    job_description: str

class PDFExportRequest(BaseModel):
    html: str


# ─── Text Extraction ─────────────────────────────────────────────────────────
def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        
        if doc.page_count > 5:
            doc.close()
            raise HTTPException(413, f"PDF has too many pages ({doc.page_count}). Please upload a resume with 5 pages or less.")
            
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
        return text.strip()
    except ImportError:
        raise HTTPException(500, "PyMuPDF not installed. Run: pip install PyMuPDF")
    except Exception as exc:
        raise HTTPException(500, f"PDF extraction failed: {exc}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        import io
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return text.strip()
    except ImportError:
        raise HTTPException(500, "python-docx not installed. Run: pip install python-docx")
    except Exception as exc:
        raise HTTPException(500, f"DOCX extraction failed: {exc}")


# ─── Contact Extraction Helpers ──────────────────────────────────────────────
def extract_email(raw_text: str) -> str | None:
    """Return the first email address found in raw text, or None."""
    matches = EMAIL_RE.findall(raw_text)
    return matches[0].lower() if matches else None


def extract_phone(raw_text: str) -> str | None:
    """Return the first plausible phone number found in raw text, or None."""
    matches = PHONE_RE.findall(raw_text)
    # Filter out short noise matches (fewer than 7 digit chars)
    for match in matches:
        digits = re.sub(r"\D", "", match)
        if len(digits) >= 7:
            return match.strip()
    return None


def hash_identifier(identifier: str) -> str:
    """Return a SHA-256 hash of the identifier for anonymized storage."""
    if not identifier:
        return ""
    return hashlib.sha256(identifier.lower().strip().encode()).hexdigest()


# ─── Supabase Cache Logic ────────────────────────────────────────────────────
def check_rate_limit(email: str) -> None:
    """
    Raise HTTP 429 if this email (hashed) was processed within the last TTL_HOURS.
    """
    sb = get_supabase()
    hashed_email = hash_identifier(email)
    cutoff = (datetime.now(timezone.utc) - timedelta(hours=TTL_HOURS)).isoformat()

    result = (
        sb.table("cv_modernizer_cache")
        .select("updated_at")
        .eq("email", hashed_email)
        .gte("updated_at", cutoff)
        .execute()
    )

    if result.data:
        # Parse updated_at and compute how long until the TTL expires
        raw_ts = result.data[0]["updated_at"]
        # Supabase returns UTC with +00:00 suffix — parse it safely
        updated_at = datetime.fromisoformat(raw_ts.replace("Z", "+00:00"))
        expires_at = updated_at + timedelta(hours=TTL_HOURS)
        remaining = expires_at - datetime.now(timezone.utc)
        total_secs = max(int(remaining.total_seconds()), 0)
        hrs = total_secs // 3600
        mins = (total_secs % 3600) // 60

        retry_msg = (
            f"{hrs}h {mins}m" if hrs else f"{mins} minute(s)"
        )
        raise HTTPException(
            status_code=429,
            detail=(
                f"This resume has already been processed recently. "
                f"Please try again in {retry_msg}."
            ),
        )


def upsert_cache(
    email: str,
    phone: str,
    experiences: list,
    education: list,
) -> None:
    """
    Explicit insert-or-update into cv_cache using hashed identifiers.
    We do NOT store names or raw PII in the database.
    """
    sb = get_supabase()
    hashed_email = hash_identifier(email)
    hashed_phone = hash_identifier(phone)
    now = datetime.now(timezone.utc).isoformat()

    # Strip bullets and skills — store only 1-liner experience metadata
    minimal_experiences = [
        {
            "company":   exp.get("company", ""),
            "role":      exp.get("role", ""),
            "startDate": exp.get("startDate", ""),
            "endDate":   exp.get("endDate", ""),
        }
        for exp in experiences
    ]

    minimal_education = [
        {
            "institution": edu.get("institution", ""),
            "degree":      edu.get("degree", ""),
            "startDate":   edu.get("startDate", ""),
            "endDate":     edu.get("endDate", ""),
        }
        for edu in education
    ]

    try:
        # Check if row exists for this hashed email
        existing = (
            sb.table("cv_modernizer_cache")
            .select("id")
            .eq("email", hashed_email)
            .execute()
        )

        if existing.data:
            sb.table("cv_modernizer_cache").update({
                "phone":       hashed_phone,
                "experiences": minimal_experiences,
                "education":   minimal_education,
                "updated_at":  now,
            }).eq("email", hashed_email).execute()
            log.info(f"Cache UPDATE for anonymized email hash={hashed_email[:8]}...")
        else:
            sb.table("cv_modernizer_cache").insert({
                "email":       hashed_email,
                "phone":       hashed_phone,
                "experiences": minimal_experiences,
                "education":   minimal_education,
                "updated_at":  now,
            }).execute()
            log.info(f"Cache INSERT for anonymized email hash={hashed_email[:8]}...")

    except Exception as exc:
        # Cache failure is non-fatal — log it but don't block the user
        log.error(f"Supabase cache write failed for hashed_email={hashed_email[:8]}...: {exc}")


# ─── Gemini Parsing ──────────────────────────────────────────────────────────
RESUME_JSON_SCHEMA = """{
  "bio": {
    "name": "string",
    "title": "string (current/most recent job title)",
    "email": "string",
    "phone": "string",
    "location": "string (City, Country)",
    "summary": "string (professional summary, 2-4 sentences)"
  },
  "links": [{"label": "string (e.g. GitHub, LinkedIn, Portfolio)", "url": "string"}],
  "experience": [
    {
      "id": "exp1",
      "company": "string",
      "role": "string",
      "startDate": "string (e.g. Jan 2021)",
      "endDate": "string (e.g. Present or Dec 2023)",
      "location": "string",
      "bullets": ["string (achievement-focused bullet point)"]
    }
  ],
  "education": [
    {
      "id": "edu1",
      "institution": "string",
      "degree": "string",
      "startDate": "string (year)",
      "endDate": "string (year)"
    }
  ],
  "skills": [
    {"category": "string (e.g. Languages, Frameworks, Tools)", "items": ["string"]}
  ]
}"""


def extract_json_from_text(text: str) -> dict:
    """
    Robustly extract a JSON object from LLM output that may contain:
    - markdown fences (```json ... ``` or ``` ... ```)
    - leading/trailing explanation text
    - extra whitespace
    Strategy: find the first '{' and last '}' and parse that slice.
    """
    text = text.strip()

    # Remove markdown fences (handles multi-line, with or without language tag)
    text = re.sub(r"```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"```", "", text)
    text = text.strip()

    # Try parsing directly first
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Fallback: find outermost { ... } block
    start = text.find("{")
    end   = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        try:
            return json.loads(text[start : end + 1])
        except json.JSONDecodeError as exc:
            raise ValueError(f"Could not parse JSON from LLM response: {exc}\n\nRaw (trimmed):\n{text[:500]}")

    raise ValueError(f"No JSON object found in LLM response.\n\nRaw (trimmed):\n{text[:500]}")


def parse_with_gemini(raw_text: str) -> dict:
    api_key = os.environ.get("GEMINI_API_KEY", "").strip()
    if not api_key:
        raise HTTPException(503, "GEMINI_API_KEY not set. Add it to your .env file.")

    try:
        from google import genai
        client = genai.Client(api_key=api_key)
    except Exception as exc:
        raise HTTPException(500, f"Gemini SDK init failed: {exc}")

    prompt = f"""You are a precise resume parser. Extract all information from the resume text and return it as valid JSON matching the schema exactly.

RULES:
- Assign sequential IDs: exp1, exp2 … for experience; edu1, edu2 … for education
- Rewrite experience bullets to be achievement-oriented if they aren't already
- Group skills logically: Languages, Frameworks, Tools, Cloud, etc.
- If a field is missing or unknown, use an empty string ""
- Return ONLY the JSON object — no markdown fences, no explanation, no extra text

JSON Schema:
{RESUME_JSON_SCHEMA}

Resume Text:
{raw_text}

JSON:"""

    log.info(f"Calling {GEMINI_MODEL}…")
    try:
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt,
        )
        raw_response = response.text
        log.info(f"Gemini responded ({len(raw_response)} chars)")
    except Exception as exc:
        err_str = str(exc)
        if "429" in err_str or "quota" in err_str.lower() or "rate" in err_str.lower():
            raise HTTPException(
                503,
                "Gemini API quota exceeded. Please wait a moment and try again."
            )
        raise HTTPException(500, f"Gemini API call failed: {exc}")

    try:
        parsed = extract_json_from_text(raw_response)
        log.info("JSON extraction from Gemini response succeeded")
        return parsed
    except ValueError as exc:
        log.error(f"JSON extraction failed: {exc}")
        raise HTTPException(500, str(exc))


# ─── API Routes ───────────────────────────────────────────────────────────────
@app.get("/api/health")
async def health_check():
    """Simple liveness probe — also reports which env vars are set."""
    return {
        "status":           "ok",
        "gemini_key_set":   bool(os.environ.get("GEMINI_API_KEY")),
        "supabase_url_set": bool(os.environ.get("SUPABASE_URL")),
        "supabase_key_set": bool(os.environ.get("SUPABASE_KEY")),
    }


@app.post("/api/extract")
@limiter.limit("10/hour")
async def extract_resume(request: Request, file: UploadFile = File(...)):
    """
    Full extraction pipeline:
      File → Text → Guard (email+phone) → Rate-limit → Gemini → Cache → Return
    """
    # ── 0. Validate file type ────────────────────────────────────────────────
    filename = (file.filename or "").lower()
    if not (filename.endswith(".pdf") or filename.endswith(".docx")):
        raise HTTPException(400, "Only PDF and DOCX files are accepted.")

    # ── 1. Read bytes in-memory (no disk write) ──────────────────────────────
    file_bytes = await file.read()
    log.info(f"Received file: {file.filename!r}  size={len(file_bytes):,} bytes")

    if len(file_bytes) > MAX_FILE_BYTES:
        raise HTTPException(
            413,
            f"File too large ({len(file_bytes) / 1_048_576:.1f} MB). "
            f"Maximum allowed size is {MAX_FILE_BYTES // 1_048_576} MB. "
            "Please compress or re-export your resume as a smaller PDF."
        )

    # ── 2. Extract raw text ──────────────────────────────────────────────────
    raw_text = (
        extract_text_from_pdf(file_bytes)
        if filename.endswith(".pdf")
        else extract_text_from_docx(file_bytes)
    )

    if not raw_text or len(raw_text) < 50:
        raise HTTPException(422, "Could not extract meaningful text from the file.")

    if len(raw_text) > MAX_TEXT_CHARS:
        log.warning(
            f"Extracted text ({len(raw_text):,} chars) exceeds limit; "
            f"truncating to {MAX_TEXT_CHARS:,} chars before Gemini call."
        )
        raw_text = raw_text[:MAX_TEXT_CHARS]

    log.info(f"Text ready: {len(raw_text):,} chars")

    # ── 3. Guard: must contain email AND phone ───────────────────────────────
    email = extract_email(raw_text)
    phone = extract_phone(raw_text)

    if not email:
        raise HTTPException(
            422,
            "No email address found in the document. "
            "Please ensure your resume contains a valid email before uploading.",
        )

    if not phone:
        raise HTTPException(
            422,
            "No phone number found in the document. "
            "Please ensure your resume contains a phone number before uploading.",
        )

    log.info(f"Detected PII (will be hashed). Email found: {'yes' if email else 'no'} | Phone found: {'yes' if phone else 'no'}")

    # ── 4. Rate-limit check via Supabase ─────────────────────────────────────
    check_rate_limit(email)   # raises HTTP 429 if within TTL window

    # ── 5. Parse with Gemini ─────────────────────────────────────────────────
    log.info("Step 5: Calling Gemini…")
    resume_data = parse_with_gemini(raw_text)
    log.info("Step 5 complete: Gemini parsing succeeded")

    # ── 6. Cache minimal record to Supabase (hashed IDs, no bullets, no skills)
    log.info("Step 6: Writing to Supabase cache (anonymized)…")
    upsert_cache(
        email=email,
        phone=phone,
        experiences=resume_data.get("experience", []),
        education=resume_data.get("education", []),
    )
    log.info("Step 6 complete: cache write done")

    # ── 7. Return full structured resume to frontend ─────────────────────────
    return {
        "resume_data":      resume_data,
        "detected_email":   email,
        "detected_phone":   phone,
        "raw_text_length":  len(raw_text),
    }


@app.post("/api/tailor")
async def tailor_resume(request: GapAnalysisRequest):
    """Use Gemini to perform gap analysis and suggest improvements against a JD."""
    api_key = os.environ.get("GEMINI_API_KEY", "").strip()
    if not api_key:
        raise HTTPException(503, "GEMINI_API_KEY not set.")

    try:
        from google import genai
        client = genai.Client(api_key=api_key)

        resume_json = json.dumps(request.resume_data.model_dump(), indent=2)

        prompt = f"""You are an expert career coach. Analyze the resume against the job description and return JSON:
{{
  "missing_skills":    ["skill1", "skill2"],
  "suggestions":       "Plain text paragraph of general suggestions",
  "improved_summary":  "Rewritten summary aligned to the JD",
  "improved_bullets":  [{{"exp_id": "exp1", "bullet_idx": 0, "new_text": "..."}}]
}}

Only return the JSON object — no extra text.

Job Description:
{request.job_description}

Current Resume:
{resume_json}

JSON:"""

        response = client.models.generate_content(
            model="gemini-1.5-flash",
            contents=prompt,
        )
        return {"analysis": extract_json_from_text(response.text)}

    except ValueError as exc:
        raise HTTPException(500, f"Gemini returned malformed JSON: {exc}")
    except Exception as exc:
        raise HTTPException(500, str(exc))


import threading
# Limit concurrent Playwright processes to prevent RAM exhaustion on VPS
pdf_semaphore = threading.Semaphore(3)

@app.post("/api/export/pdf")
@limiter.limit("20/hour")
def export_pdf(request: Request, req: PDFExportRequest):
    """
    Generate a PDF from raw HTML using Playwright.
    Runs synchronously in a threadpool to avoid Windows asyncio loop conflicts.
    Protected by a Semaphore to limit concurrent browser spawns.
    """
    if not pdf_semaphore.acquire(timeout=30.0):
        raise HTTPException(503, "Server is currently generating too many PDFs. Please try again in a few seconds.")
        
    try:
        from playwright.sync_api import sync_playwright
        import io
        from fastapi.responses import StreamingResponse

        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            
            # Use waitUntil="networkidle" to ensure fonts load
            page.set_content(req.html, wait_until="networkidle")
            
            # Evaluate the exact height of the rendered resume content
            height_px = page.evaluate("""() => {
                const sheet = document.querySelector('.resume-sheet');
                return sheet ? sheet.getBoundingClientRect().height : document.body.scrollHeight;
            }""")
            
            pdf_bytes = page.pdf(
                width="210mm",
                height=f"{height_px}px",
                print_background=True,
                margin={"top": "0", "right": "0", "bottom": "0", "left": "0"}
            )
            browser.close()

        buf = io.BytesIO(pdf_bytes)
        buf.seek(0)
        
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": 'attachment; filename="resume.pdf"'},
        )
    except ImportError:
        log.error("Playwright not installed.")
        raise HTTPException(500, "Playwright is not installed on the server.")
    except Exception as exc:
        log.error(f"PDF export failed: {exc}")
        raise HTTPException(500, f"PDF generation failed: {exc}")
    finally:
        pdf_semaphore.release()


@app.post("/api/export/docx")
@limiter.limit("20/hour")
async def export_docx(request: Request, data: ResumeData):
    """
    Generate a clean Word document from the resume data.
    Returns the .docx as a binary stream for direct browser download.
    """
    try:
        import io
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from fastapi.responses import StreamingResponse

        doc = Document()

        # ── Page margins ────────────────────────────────────────────────
        for section in doc.sections:
            section.top_margin    = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin   = Inches(0.85)
            section.right_margin  = Inches(0.85)

        bio = data.bio or {}

        # ── Helper styles ───────────────────────────────────────────────
        def set_font(run, size_pt: float, bold=False, color: tuple | None = None):
            run.font.name = "Calibri"
            run.font.size = Pt(size_pt)
            run.font.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*color)

        def add_section_heading(title: str):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(title.upper())
            set_font(run, 10, bold=True, color=(70, 90, 160))
            # Underline via bottom border
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"),  "single")
            bottom.set(qn("w:sz"),   "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "4a5aa0")
            pBdr.append(bottom)
            pPr.append(pBdr)
            return p

        def add_thin_line(para):
            """No-op — section heading already has a border."""
            pass

        # ── Name ────────────────────────────────────────────────────────
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(bio.get("name", "Your Name"))
        set_font(name_run, 22, bold=True, color=(30, 30, 30))

        # ── Job title ───────────────────────────────────────────────────
        if bio.get("title"):
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_before = Pt(1)
            t_run = title_para.add_run(bio["title"])
            set_font(t_run, 12, color=(80, 80, 80))

        # ── Contact line ─────────────────────────────────────────────────
        contact_parts = [v for k in ("email", "phone", "location") if (v := bio.get(k))]
        if contact_parts:
            c_para = doc.add_paragraph()
            c_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c_para.paragraph_format.space_before = Pt(2)
            c_run = c_para.add_run("  ·  ".join(contact_parts))
            set_font(c_run, 9, color=(100, 100, 100))

        # ── Links ────────────────────────────────────────────────────────
        if data.links:
            link_labels = [lnk.get("label") or lnk.get("url", "") for lnk in data.links if lnk.get("url")]
            if link_labels:
                lp = doc.add_paragraph()
                lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                lp.paragraph_format.space_before = Pt(1)
                lr = lp.add_run("  |  ".join(link_labels))
                set_font(lr, 9, color=(70, 90, 160))

        # ── Summary ──────────────────────────────────────────────────────
        if bio.get("summary"):
            add_section_heading("Professional Summary")
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(2)
            sr = sp.add_run(bio["summary"])
            set_font(sr, 10)

        # ── Experience ───────────────────────────────────────────────────
        if data.experience:
            add_section_heading("Experience")
            for exp in data.experience:
                # Role · Company
                ep = doc.add_paragraph()
                ep.paragraph_format.space_before = Pt(6)
                ep.paragraph_format.space_after  = Pt(0)
                role_run = ep.add_run(exp.get("role", ""))
                set_font(role_run, 11, bold=True)
                company = exp.get("company", "")
                location = exp.get("location", "")
                if company:
                    company_str = f"  ·  {company}"
                    if location:
                        company_str += f", {location}"
                    c_run = ep.add_run(company_str)
                    set_font(c_run, 10, color=(80, 80, 80))
                # Date range (right-aligned via tab)
                start = exp.get("startDate", "")
                end   = exp.get("endDate",   "")
                if start or end:
                    dp = doc.add_paragraph()
                    dp.paragraph_format.space_before = Pt(0)
                    dp.paragraph_format.space_after  = Pt(2)
                    d_run = dp.add_run(f"{start} – {end}")
                    set_font(d_run, 9, color=(120, 120, 120))
                # Bullets
                for bullet in exp.get("bullets", []):
                    if bullet.strip():
                        bp = doc.add_paragraph(style="List Bullet")
                        bp.paragraph_format.space_before = Pt(1)
                        bp.paragraph_format.space_after  = Pt(1)
                        bp.paragraph_format.left_indent  = Inches(0.2)
                        b_run = bp.add_run(bullet)
                        set_font(b_run, 10)

        # ── Education ────────────────────────────────────────────────────
        if data.education:
            add_section_heading("Education")
            for edu in data.education:
                edp = doc.add_paragraph()
                edp.paragraph_format.space_before = Pt(5)
                edp.paragraph_format.space_after  = Pt(0)
                deg_run = edp.add_run(edu.get("degree", ""))
                set_font(deg_run, 11, bold=True)
                inst = edu.get("institution", "")
                if inst:
                    i_run = edp.add_run(f"  ·  {inst}")
                    set_font(i_run, 10, color=(80, 80, 80))
                start = edu.get("startDate", "")
                end   = edu.get("endDate",   "")
                if start or end:
                    dp = doc.add_paragraph()
                    dp.paragraph_format.space_before = Pt(0)
                    dp.paragraph_format.space_after  = Pt(2)
                    d_run = dp.add_run(f"{start} – {end}")
                    set_font(d_run, 9, color=(120, 120, 120))

        # ── Skills ───────────────────────────────────────────────────────
        if data.skills:
            add_section_heading("Skills")
            for group in data.skills:
                category = group.get("category", "")
                items    = group.get("items", [])
                if not items:
                    continue
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(3)
                sp.paragraph_format.space_after  = Pt(1)
                if category:
                    cat_run = sp.add_run(f"{category}: ")
                    set_font(cat_run, 10, bold=True)
                items_run = sp.add_run(", ".join(i for i in items if i.strip()))
                set_font(items_run, 10)

        # ── Serialise to bytes ───────────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        safe_name = re.sub(r"[^\w\-]", "_", bio.get("name", "resume")).strip("_") or "resume"
        filename  = f"{safe_name}_resume.docx"

        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    except ImportError as exc:
        raise HTTPException(500, f"python-docx not installed: {exc}")
    except Exception as exc:
        log.error(f"DOCX generation failed: {exc}")
        raise HTTPException(500, f"DOCX generation failed: {exc}")


# ─── Entry point ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)

